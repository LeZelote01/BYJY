#!/usr/bin/env python3
"""
Extracteur de Métadonnées
========================

Outil forensique pour extraire les métadonnées de fichiers multiples formats.
Support images (EXIF), documents (PDF, Office), audio/vidéo et autres formats.

Auteur: Jean Yves (LeZelote)
Date: Mai 2025
Version: 1.0
"""

import os
import sys
import argparse
import json
import csv
from datetime import datetime
from pathlib import Path
import hashlib
import mimetypes

# Import des bibliothèques de métadonnées
try:
    from PIL import Image
    from PIL.ExifTags import TAGS, GPSTAGS
    PILLOW_AVAILABLE = True
except ImportError:
    PILLOW_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from mutagen import File as MutagenFile
    from mutagen.id3 import ID3NoHeaderError
    MUTAGEN_AVAILABLE = True
except ImportError:
    MUTAGEN_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# import tkinter as tk
# from tkinter import ttk, filedialog, messagebox, scrolledtext
# import threading


class MetadataExtractor:
    """Classe principale pour l'extraction de métadonnées."""
    
    def __init__(self):
        self.supported_formats = {
            'images': ['.jpg', '.jpeg', '.png', '.tiff', '.tif', '.bmp', '.gif'],
            'documents': ['.pdf', '.docx', '.doc', '.txt'],
            'audio': ['.mp3', '.wav', '.flac', '.ogg', '.m4a', '.aac'],
            'video': ['.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv'],
            'archives': ['.zip', '.rar', '.7z', '.tar', '.gz']
        }
        
        self.extraction_methods = {
            'image': self._extract_image_metadata,
            'pdf': self._extract_pdf_metadata,
            'audio': self._extract_audio_metadata,
            'video': self._extract_video_metadata,
            'document': self._extract_document_metadata,
            'generic': self._extract_generic_metadata
        }
    
    def detect_file_type(self, file_path):
        """
        Détecte le type de fichier basé sur son extension et MIME type.
        
        Args:
            file_path (str): Chemin vers le fichier
        
        Returns:
            str: Type de fichier détecté
        """
        file_ext = Path(file_path).suffix.lower()
        mime_type, _ = mimetypes.guess_type(file_path)
        
        # Vérification par extension
        for category, extensions in self.supported_formats.items():
            if file_ext in extensions:
                if category == 'images':
                    return 'image'
                elif category == 'documents':
                    if file_ext == '.pdf':
                        return 'pdf'
                    else:
                        return 'document'
                elif category == 'audio':
                    return 'audio'
                elif category == 'video':
                    return 'video'
        
        # Vérification par MIME type
        if mime_type:
            if mime_type.startswith('image/'):
                return 'image'
            elif mime_type.startswith('audio/'):
                return 'audio'
            elif mime_type.startswith('video/'):
                return 'video'
            elif mime_type == 'application/pdf':
                return 'pdf'
        
        return 'generic'
    
    def extract_metadata(self, file_path, include_hash=True, progress_callback=None):
        """
        Extrait les métadonnées d'un fichier.
        
        Args:
            file_path (str): Chemin vers le fichier
            include_hash (bool): Inclure les hachages du fichier
            progress_callback (function): Callback pour le progrès
        
        Returns:
            dict: Métadonnées extraites
        """
        try:
            if not os.path.exists(file_path):
                return {'error': 'Fichier inexistant'}
            
            if progress_callback:
                progress_callback(f"Analyse de {os.path.basename(file_path)}")
            
            # Métadonnées de base du fichier
            file_stats = os.stat(file_path)
            metadata = {
                'file_info': {
                    'filename': os.path.basename(file_path),
                    'filepath': os.path.abspath(file_path),
                    'size_bytes': file_stats.st_size,
                    'size_human': self._format_file_size(file_stats.st_size),
                    'created_time': datetime.fromtimestamp(file_stats.st_ctime).isoformat(),
                    'modified_time': datetime.fromtimestamp(file_stats.st_mtime).isoformat(),
                    'accessed_time': datetime.fromtimestamp(file_stats.st_atime).isoformat(),
                    'extension': Path(file_path).suffix.lower(),
                    'mime_type': mimetypes.guess_type(file_path)[0]
                },
                'extraction_info': {
                    'extracted_at': datetime.now().isoformat(),
                    'extractor_version': '1.0',
                    'file_type_detected': self.detect_file_type(file_path)
                }
            }
            
            # Calculer les hachages si demandé
            if include_hash:
                metadata['hashes'] = self._calculate_file_hashes(file_path, progress_callback)
            
            # Extraction spécialisée selon le type
            file_type = metadata['extraction_info']['file_type_detected']
            extraction_method = self.extraction_methods.get(file_type, self.extraction_methods['generic'])
            
            specialized_metadata = extraction_method(file_path)
            if specialized_metadata:
                metadata.update(specialized_metadata)
            
            return metadata
            
        except Exception as e:
            return {
                'error': str(e),
                'file_info': {
                    'filename': os.path.basename(file_path),
                    'filepath': os.path.abspath(file_path)
                }
            }
    
    def _calculate_file_hashes(self, file_path, progress_callback=None):
        """Calcule plusieurs hachages du fichier."""
        hashes = {}
        algorithms = ['md5', 'sha1', 'sha256']
        
        hash_objects = {alg: hashlib.new(alg) for alg in algorithms}
        
        try:
            with open(file_path, 'rb') as f:
                while chunk := f.read(8192):
                    for hash_obj in hash_objects.values():
                        hash_obj.update(chunk)
                    
                    if progress_callback:
                        progress_callback(f"Calcul des hachages...")
            
            for alg, hash_obj in hash_objects.items():
                hashes[alg] = hash_obj.hexdigest()
            
        except Exception as e:
            hashes['error'] = str(e)
        
        return hashes
    
    def _extract_image_metadata(self, file_path):
        """Extrait les métadonnées des images (EXIF)."""
        if not PILLOW_AVAILABLE:
            return {'image_metadata': {'error': 'Pillow non disponible'}}
        
        try:
            with Image.open(file_path) as image:
                metadata = {
                    'image_metadata': {
                        'format': image.format,
                        'mode': image.mode,
                        'size': image.size,
                        'width': image.size[0],
                        'height': image.size[1],
                        'has_transparency': hasattr(image, 'transparency') and image.transparency is not None
                    }
                }
                
                # Extraction des données EXIF
                if hasattr(image, '_getexif'):
                    exif_data = image._getexif()
                    if exif_data:
                        exif_metadata = {}
                        gps_metadata = {}
                        
                        for tag_id, value in exif_data.items():
                            tag = TAGS.get(tag_id, tag_id)
                            
                            # Traitement spécial pour les données GPS
                            if tag == 'GPSInfo':
                                for gps_tag_id, gps_value in value.items():
                                    gps_tag = GPSTAGS.get(gps_tag_id, gps_tag_id)
                                    gps_metadata[gps_tag] = str(gps_value)
                            else:
                                # Convertir en string pour éviter les erreurs de sérialisation
                                try:
                                    exif_metadata[tag] = str(value)
                                except:
                                    exif_metadata[tag] = '<non-serializable>'
                        
                        metadata['image_metadata']['exif'] = exif_metadata
                        if gps_metadata:
                            metadata['image_metadata']['gps'] = gps_metadata
                            # Essayer d'extraire les coordonnées lisibles
                            coords = self._parse_gps_coordinates(gps_metadata)
                            if coords:
                                metadata['image_metadata']['coordinates'] = coords
                
                return metadata
                
        except Exception as e:
            return {'image_metadata': {'error': str(e)}}
    
    def _parse_gps_coordinates(self, gps_data):
        """Parse les coordonnées GPS au format lisible."""
        try:
            lat_ref = gps_data.get('GPSLatitudeRef', '')
            lat_data = gps_data.get('GPSLatitude', '')
            lon_ref = gps_data.get('GPSLongitudeRef', '')
            lon_data = gps_data.get('GPSLongitude', '')
            
            if lat_data and lon_data:
                # Conversion simple (format peut varier)
                return {
                    'latitude': f"{lat_data} {lat_ref}",
                    'longitude': f"{lon_data} {lon_ref}",
                    'note': 'Format brut - conversion manuelle nécessaire'
                }
        except:
            pass
        return None
    
    def _extract_pdf_metadata(self, file_path):
        """Extrait les métadonnées des fichiers PDF."""
        if not PYPDF2_AVAILABLE:
            return {'pdf_metadata': {'error': 'PyPDF2 non disponible'}}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    'pdf_metadata': {
                        'pages_count': len(pdf_reader.pages),
                        'is_encrypted': pdf_reader.is_encrypted
                    }
                }
                
                # Métadonnées du document
                if pdf_reader.metadata:
                    doc_info = {}
                    for key, value in pdf_reader.metadata.items():
                        # Nettoyer les clés (enlever le préfixe /)
                        clean_key = key.replace('/', '') if key.startswith('/') else key
                        doc_info[clean_key] = str(value) if value else ''
                    
                    metadata['pdf_metadata']['document_info'] = doc_info
                
                # Informations sur la première page (si accessible)
                if not pdf_reader.is_encrypted and len(pdf_reader.pages) > 0:
                    first_page = pdf_reader.pages[0]
                    if hasattr(first_page, 'mediabox'):
                        mediabox = first_page.mediabox
                        metadata['pdf_metadata']['page_size'] = {
                            'width': float(mediabox.width),
                            'height': float(mediabox.height),
                            'unit': 'points'
                        }
                
                return metadata
                
        except Exception as e:
            return {'pdf_metadata': {'error': str(e)}}
    
    def _extract_audio_metadata(self, file_path):
        """Extrait les métadonnées des fichiers audio."""
        if not MUTAGEN_AVAILABLE:
            return {'audio_metadata': {'error': 'Mutagen non disponible'}}
        
        try:
            audio_file = MutagenFile(file_path)
            
            if audio_file is None:
                return {'audio_metadata': {'error': 'Format audio non supporté'}}
            
            metadata = {
                'audio_metadata': {
                    'length_seconds': getattr(audio_file.info, 'length', 0),
                    'bitrate': getattr(audio_file.info, 'bitrate', 0),
                    'sample_rate': getattr(audio_file.info, 'sample_rate', 0),
                    'channels': getattr(audio_file.info, 'channels', 0),
                    'format': audio_file.mime[0] if audio_file.mime else 'unknown'
                }
            }
            
            # Durée au format lisible
            length = metadata['audio_metadata']['length_seconds']
            if length > 0:
                minutes, seconds = divmod(int(length), 60)
                hours, minutes = divmod(minutes, 60)
                if hours > 0:
                    duration = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
                else:
                    duration = f"{minutes:02d}:{seconds:02d}"
                metadata['audio_metadata']['duration'] = duration
            
            # Tags du fichier
            if audio_file.tags:
                tags = {}
                for key, value in audio_file.tags.items():
                    # Convertir les valeurs en string
                    if isinstance(value, list):
                        tags[key] = [str(v) for v in value]
                    else:
                        tags[key] = str(value)
                
                metadata['audio_metadata']['tags'] = tags
                
                # Tags courants avec noms lisibles
                common_tags = {
                    'TIT2': 'title',        # MP3
                    'TPE1': 'artist',       # MP3
                    'TALB': 'album',        # MP3
                    'TDRC': 'date',         # MP3
                    'TCON': 'genre',        # MP3
                    'TITLE': 'title',       # Autres formats
                    'ARTIST': 'artist',     # Autres formats
                    'ALBUM': 'album',       # Autres formats
                    'DATE': 'date',         # Autres formats
                    'GENRE': 'genre'        # Autres formats
                }
                
                readable_tags = {}
                for tag_key, tag_value in tags.items():
                    readable_key = common_tags.get(tag_key.upper(), tag_key)
                    readable_tags[readable_key] = tag_value
                
                metadata['audio_metadata']['readable_tags'] = readable_tags
            
            return metadata
            
        except ID3NoHeaderError:
            return {'audio_metadata': {'error': 'Pas de métadonnées ID3'}}
        except Exception as e:
            return {'audio_metadata': {'error': str(e)}}
    
    def _extract_video_metadata(self, file_path):
        """Extrait les métadonnées des fichiers vidéo (basique)."""
        # Note: Pour une extraction complète, ffprobe serait idéal
        # Ici on utilise les capacités de base
        
        try:
            file_size = os.path.getsize(file_path)
            
            metadata = {
                'video_metadata': {
                    'container': Path(file_path).suffix.lower().replace('.', ''),
                    'file_size_mb': round(file_size / (1024 * 1024), 2),
                    'note': 'Métadonnées basiques - ffprobe recommandé pour plus de détails'
                }
            }
            
            # Tentative d'extraction avec mutagen (pour certains formats)
            if MUTAGEN_AVAILABLE:
                try:
                    video_file = MutagenFile(file_path)
                    if video_file and hasattr(video_file.info, 'length'):
                        length = getattr(video_file.info, 'length', 0)
                        if length > 0:
                            minutes, seconds = divmod(int(length), 60)
                            hours, minutes = divmod(minutes, 60)
                            duration = f"{hours:02d}:{minutes:02d}:{seconds:02d}"
                            metadata['video_metadata']['duration'] = duration
                            metadata['video_metadata']['length_seconds'] = length
                        
                        if hasattr(video_file.info, 'bitrate'):
                            metadata['video_metadata']['bitrate'] = video_file.info.bitrate
                        
                        if video_file.tags:
                            tags = {}
                            for key, value in video_file.tags.items():
                                tags[key] = str(value) if not isinstance(value, list) else [str(v) for v in value]
                            metadata['video_metadata']['tags'] = tags
                            
                except:
                    pass
            
            return metadata
            
        except Exception as e:
            return {'video_metadata': {'error': str(e)}}
    
    def _extract_document_metadata(self, file_path):
        """Extrait les métadonnées des documents Office."""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.docx' and DOCX_AVAILABLE:
            return self._extract_docx_metadata(file_path)
        elif file_ext == '.txt':
            return self._extract_text_metadata(file_path)
        else:
            return {'document_metadata': {'error': f'Format {file_ext} non supporté'}}
    
    def _extract_docx_metadata(self, file_path):
        """Extrait les métadonnées des fichiers DOCX."""
        try:
            doc = DocxDocument(file_path)
            
            metadata = {
                'document_metadata': {
                    'type': 'Microsoft Word Document',
                    'paragraphs_count': len(doc.paragraphs),
                    'tables_count': len(doc.tables)
                }
            }
            
            # Propriétés du document
            if doc.core_properties:
                core_props = {}
                props_to_extract = [
                    'title', 'author', 'subject', 'keywords', 'comments',
                    'created', 'modified', 'last_modified_by', 'revision',
                    'version', 'category', 'content_status'
                ]
                
                for prop in props_to_extract:
                    try:
                        value = getattr(doc.core_properties, prop, None)
                        if value is not None:
                            # Convertir les datetime en string
                            if hasattr(value, 'isoformat'):
                                core_props[prop] = value.isoformat()
                            else:
                                core_props[prop] = str(value)
                    except:
                        pass
                
                if core_props:
                    metadata['document_metadata']['core_properties'] = core_props
            
            # Statistiques du texte
            text_content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            if text_content:
                metadata['document_metadata']['text_stats'] = {
                    'characters_count': len(text_content),
                    'characters_count_no_spaces': len(text_content.replace(' ', '')),
                    'words_count': len(text_content.split()),
                    'lines_count': len(text_content.split('\n'))
                }
            
            return metadata
            
        except Exception as e:
            return {'document_metadata': {'error': str(e)}}
    
    def _extract_text_metadata(self, file_path):
        """Extrait les métadonnées des fichiers texte."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            metadata = {
                'document_metadata': {
                    'type': 'Plain Text',
                    'encoding': 'UTF-8 (détecté)',
                    'characters_count': len(content),
                    'characters_count_no_spaces': len(content.replace(' ', '')),
                    'words_count': len(content.split()),
                    'lines_count': len(content.split('\n')),
                    'paragraphs_count': len([p for p in content.split('\n\n') if p.strip()])
                }
            }
            
            # Analyse basique du contenu
            if content:
                first_line = content.split('\n')[0][:100]
                metadata['document_metadata']['first_line_preview'] = first_line
                
                # Détection de possibles formats structurés
                if content.strip().startswith('{') and content.strip().endswith('}'):
                    metadata['document_metadata']['possible_format'] = 'JSON'
                elif content.strip().startswith('<') and content.strip().endswith('>'):
                    metadata['document_metadata']['possible_format'] = 'XML/HTML'
                elif ',' in content and '\n' in content:
                    metadata['document_metadata']['possible_format'] = 'CSV'
            
            return metadata
            
        except Exception as e:
            return {'document_metadata': {'error': str(e)}}
    
    def _extract_generic_metadata(self, file_path):
        """Extraction générique pour tous types de fichiers."""
        try:
            # Analyse des signatures de fichier (magic numbers)
            with open(file_path, 'rb') as f:
                header = f.read(16)
            
            metadata = {
                'generic_metadata': {
                    'header_hex': header.hex(),
                    'header_ascii': ''.join(chr(b) if 32 <= b <= 126 else '.' for b in header)
                }
            }
            
            # Détection de formats basée sur les signatures
            signatures = {
                b'\x50\x4B\x03\x04': 'ZIP Archive',
                b'\x50\x4B\x05\x06': 'ZIP Archive (empty)',
                b'\x50\x4B\x07\x08': 'ZIP Archive (spanned)',
                b'\x52\x61\x72\x21': 'RAR Archive',
                b'\x37\x7A\xBC\xAF': '7-Zip Archive',
                b'\x1F\x8B': 'GZIP Archive',
                b'\x42\x5A\x68': 'BZIP2 Archive',
                b'\x89\x50\x4E\x47': 'PNG Image',
                b'\xFF\xD8\xFF': 'JPEG Image',
                b'\x47\x49\x46\x38': 'GIF Image',
                b'\x25\x50\x44\x46': 'PDF Document',
                b'\xD0\xCF\x11\xE0': 'Microsoft Office Document',
                b'\x50\x4B\x03\x04': 'Microsoft Office (ZIP-based)',
            }
            
            for signature, file_type in signatures.items():
                if header.startswith(signature):
                    metadata['generic_metadata']['detected_format'] = file_type
                    break
            
            return metadata
            
        except Exception as e:
            return {'generic_metadata': {'error': str(e)}}
    
    def _format_file_size(self, size_bytes):
        """Formate la taille du fichier en unités lisibles."""
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} PB"
    
    def extract_batch(self, file_list, progress_callback=None):
        """
        Extrait les métadonnées de plusieurs fichiers.
        
        Args:
            file_list (list): Liste des chemins de fichiers
            progress_callback (function): Callback pour le progrès
        
        Returns:
            list: Liste des métadonnées extraites
        """
        results = []
        total_files = len(file_list)
        
        for i, file_path in enumerate(file_list):
            if progress_callback:
                progress = int((i / total_files) * 100)
                progress_callback(progress, f"Traitement {i+1}/{total_files}: {os.path.basename(file_path)}")
            
            metadata = self.extract_metadata(file_path)
            results.append(metadata)
        
        if progress_callback:
            progress_callback(100, f"Extraction terminée: {total_files} fichiers traités")
        
        return results
    
    def export_results(self, results, output_format='json', output_file=None):
        """
        Exporte les résultats dans différents formats.
        
        Args:
            results (list): Liste des métadonnées
            output_format (str): Format de sortie ('json', 'csv', 'html')
            output_file (str): Fichier de sortie
        
        Returns:
            str: Contenu exporté
        """
        if output_format == 'json':
            content = json.dumps(results, indent=2, ensure_ascii=False, default=str)
        elif output_format == 'csv':
            content = self._generate_csv_report(results)
        elif output_format == 'html':
            content = self._generate_html_report(results)
        else:
            content = str(results)
        
        if output_file:
            try:
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(content)
                print(f"Rapport exporté vers: {output_file}")
            except Exception as e:
                print(f"Erreur d'export: {e}")
        
        return content
    
    def _generate_csv_report(self, results):
        """Génère un rapport CSV des métadonnées."""
        import io
        output = io.StringIO()
        
        if not results:
            return "Aucun résultat à exporter"
        
        # Déterminer les colonnes communes
        all_keys = set()
        for result in results:
            def extract_keys(obj, prefix=''):
                if isinstance(obj, dict):
                    for k, v in obj.items():
                        if isinstance(v, dict):
                            extract_keys(v, f"{prefix}{k}.")
                        elif isinstance(v, list):
                            all_keys.add(f"{prefix}{k}")
                        else:
                            all_keys.add(f"{prefix}{k}")
                            
            extract_keys(result)
        
        # Colonnes principales
        main_columns = [
            'file_info.filename', 'file_info.filepath', 'file_info.size_human',
            'file_info.extension', 'file_info.mime_type', 'file_info.modified_time'
        ]
        
        # Ajouter d'autres colonnes importantes
        other_columns = sorted([k for k in all_keys if k not in main_columns])
        columns = main_columns + other_columns
        
        writer = csv.writer(output)
        writer.writerow(columns)
        
        for result in results:
            def get_nested_value(obj, key):
                keys = key.split('.')
                current = obj
                for k in keys:
                    if isinstance(current, dict) and k in current:
                        current = current[k]
                    else:
                        return ''
                return str(current) if current is not None else ''
            
            row = [get_nested_value(result, col) for col in columns]
            writer.writerow(row)
        
        return output.getvalue()
    
    def _generate_html_report(self, results):
        """Génère un rapport HTML des métadonnées."""
        html = f"""
<!DOCTYPE html>
<html lang="fr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Rapport d'Extraction de Métadonnées</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; background-color: #f5f5f5; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: white; padding: 20px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 10px; margin-bottom: 20px; }}
        .file-section {{ margin: 20px 0; border: 1px solid #ddd; border-radius: 8px; overflow: hidden; }}
        .file-header {{ background: #f8f9fa; padding: 15px; border-bottom: 1px solid #ddd; }}
        .file-content {{ padding: 15px; }}
        .metadata-section {{ margin: 10px 0; }}
        .metadata-title {{ font-weight: bold; color: #495057; margin-bottom: 10px; padding: 8px; background: #e9ecef; border-radius: 4px; }}
        .metadata-table {{ width: 100%; border-collapse: collapse; margin-bottom: 15px; }}
        .metadata-table th, .metadata-table td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        .metadata-table th {{ background: #f8f9fa; font-weight: bold; }}
        .error {{ color: #dc3545; background: #f8d7da; padding: 10px; border-radius: 4px; }}
        .success {{ color: #155724; background: #d4edda; padding: 10px; border-radius: 4px; }}
        .stats {{ display: flex; gap: 20px; margin: 20px 0; }}
        .stat {{ flex: 1; text-align: center; padding: 15px; background: #e9ecef; border-radius: 8px; }}
        .stat-value {{ font-size: 24px; font-weight: bold; color: #495057; }}
        .stat-label {{ font-size: 14px; color: #6c757d; }}
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1>🔍 Rapport d'Extraction de Métadonnées</h1>
            <p>Généré le: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>
"""
        
        # Statistiques générales
        total_files = len(results)
        successful = len([r for r in results if 'error' not in r])
        errors = total_files - successful
        
        html += f"""
        <div class="stats">
            <div class="stat">
                <div class="stat-value">{total_files}</div>
                <div class="stat-label">Fichiers analysés</div>
            </div>
            <div class="stat">
                <div class="stat-value">{successful}</div>
                <div class="stat-label">Succès</div>
            </div>
            <div class="stat">
                <div class="stat-value">{errors}</div>
                <div class="stat-label">Erreurs</div>
            </div>
        </div>
"""
        
        # Détails par fichier
        for i, result in enumerate(results, 1):
            filename = result.get('file_info', {}).get('filename', 'Fichier inconnu')
            
            html += f"""
        <div class="file-section">
            <div class="file-header">
                <h3>📁 Fichier {i}: {filename}</h3>
            </div>
            <div class="file-content">
"""
            
            if 'error' in result:
                html += f'<div class="error">❌ Erreur: {result["error"]}</div>'
            else:
                html += '<div class="success">✅ Métadonnées extraites avec succès</div>'
                
                # Afficher chaque section de métadonnées
                for section_name, section_data in result.items():
                    if isinstance(section_data, dict) and section_name != 'file_info':
                        html += f'<div class="metadata-section">'
                        html += f'<div class="metadata-title">{section_name.replace("_", " ").title()}</div>'
                        html += '<table class="metadata-table">'
                        
                        for key, value in section_data.items():
                            if isinstance(value, dict):
                                # Sous-section
                                for subkey, subvalue in value.items():
                                    html += f'<tr><th>{key}.{subkey}</th><td>{str(subvalue)}</td></tr>'
                            else:
                                html += f'<tr><th>{key}</th><td>{str(value)}</td></tr>'
                        
                        html += '</table></div>'
            
            html += '</div></div>'
        
        html += """
    </div>
</body>
</html>
"""
        return html


class MetadataExtractorGUI:
    """Interface graphique pour l'extracteur de métadonnées."""
    
    def __init__(self):
        self.extractor = MetadataExtractor()
        self.selected_files = []
        self.results = []
        self.setup_gui()
    
    def setup_gui(self):
        """Configure l'interface graphique."""
        self.root = tk.Tk()
        self.root.title("🔍 Extracteur de Métadonnées v1.0")
        self.root.geometry("900x700")
        self.root.configure(bg='#f0f0f0')
        
        # Variables
        self.include_hash_var = tk.BooleanVar(value=True)
        self.output_format_var = tk.StringVar(value='json')
        
        self.create_widgets()
    
    def create_widgets(self):
        """Crée les widgets de l'interface."""
        # Frame principal avec onglets
        notebook = ttk.Notebook(self.root)
        notebook.pack(fill='both', expand=True, padx=10, pady=10)
        
        # Onglets
        self.main_frame = ttk.Frame(notebook)
        notebook.add(self.main_frame, text='🔍 Extraction')
        
        self.results_frame = ttk.Frame(notebook)
        notebook.add(self.results_frame, text='📊 Résultats')
        
        self.about_frame = ttk.Frame(notebook)
        notebook.add(self.about_frame, text='ℹ️ À propos')
        
        self.create_main_tab()
        self.create_results_tab()
        self.create_about_tab()
    
    def create_main_tab(self):
        """Crée l'onglet principal."""
        # Sélection de fichiers
        files_frame = ttk.LabelFrame(self.main_frame, text="Sélection de fichiers", padding=10)
        files_frame.pack(fill='both', expand=True, padx=10, pady=5)
        
        # Boutons
        buttons_frame = ttk.Frame(files_frame)
        buttons_frame.pack(fill='x', pady=5)
        
        ttk.Button(buttons_frame, text="📁 Ajouter fichiers", 
                  command=self.add_files).pack(side='left', padx=5)
        ttk.Button(buttons_frame, text="📂 Ajouter dossier", 
                  command=self.add_directory).pack(side='left', padx=5)
        ttk.Button(buttons_frame, text="🗑️ Effacer", 
                  command=self.clear_files).pack(side='left', padx=5)
        
        # Liste des fichiers
        self.files_listbox = tk.Listbox(files_frame, height=10)
        scrollbar = ttk.Scrollbar(files_frame, orient='vertical', command=self.files_listbox.yview)
        self.files_listbox.configure(yscrollcommand=scrollbar.set)
        
        self.files_listbox.pack(side='left', fill='both', expand=True, pady=5)
        scrollbar.pack(side='right', fill='y', pady=5)
        
        # Options
        options_frame = ttk.LabelFrame(self.main_frame, text="Options d'extraction", padding=10)
        options_frame.pack(fill='x', padx=10, pady=5)
        
        ttk.Checkbutton(options_frame, text="Calculer les hachages (MD5, SHA1, SHA256)",
                       variable=self.include_hash_var).pack(anchor='w')
        
        # Format de sortie
        format_frame = ttk.Frame(options_frame)
        format_frame.pack(fill='x', pady=5)
        
        ttk.Label(format_frame, text="Format de sortie:").pack(side='left')
        ttk.Radiobutton(format_frame, text="JSON", variable=self.output_format_var, 
                       value='json').pack(side='left', padx=10)
        ttk.Radiobutton(format_frame, text="CSV", variable=self.output_format_var, 
                       value='csv').pack(side='left', padx=10)
        ttk.Radiobutton(format_frame, text="HTML", variable=self.output_format_var, 
                       value='html').pack(side='left', padx=10)
        
        # Bouton d'extraction
        action_frame = ttk.Frame(self.main_frame)
        action_frame.pack(fill='x', padx=10, pady=10)
        
        self.extract_button = ttk.Button(action_frame, text="🚀 Extraire les métadonnées", 
                                        command=self.start_extraction)
        self.extract_button.pack()
        
        # Progression
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(self.main_frame, variable=self.progress_var, 
                                           maximum=100, length=400)
        self.progress_bar.pack(padx=10, pady=5)
        
        self.status_var = tk.StringVar(value="Prêt")
        self.status_label = ttk.Label(self.main_frame, textvariable=self.status_var)
        self.status_label.pack(pady=5)
    
    def create_results_tab(self):
        """Crée l'onglet des résultats."""
        # Boutons d'action
        buttons_frame = ttk.Frame(self.results_frame)
        buttons_frame.pack(fill='x', padx=10, pady=5)
        
        ttk.Button(buttons_frame, text="💾 Sauvegarder", 
                  command=self.save_results).pack(side='left', padx=5)
        ttk.Button(buttons_frame, text="🔄 Actualiser", 
                  command=self.refresh_results).pack(side='left', padx=5)
        ttk.Button(buttons_frame, text="🗑️ Effacer", 
                  command=self.clear_results).pack(side='left', padx=5)
        
        # Zone de résultats
        self.results_text = scrolledtext.ScrolledText(self.results_frame, height=25)
        self.results_text.pack(fill='both', expand=True, padx=10, pady=5)
    
    def create_about_tab(self):
        """Crée l'onglet à propos."""
        about_text = """
🔍 Extracteur de Métadonnées v1.0

Outil forensique pour extraire les métadonnées de fichiers multiples formats.

Formats supportés:
• Images (EXIF): JPG, PNG, TIFF, BMP, GIF
• Documents: PDF, DOCX, TXT
• Audio: MP3, WAV, FLAC, OGG, M4A
• Vidéo: MP4, AVI, MKV, MOV (basique)
• Génériques: Tous formats avec analyse des signatures

Métadonnées extraites:
• Informations fichier: taille, dates, permissions
• Hachages cryptographiques: MD5, SHA1, SHA256
• EXIF pour images: géolocalisation, appareil photo
• Propriétés documents: auteur, titre, création
• Tags audio/vidéo: artiste, album, durée
• Signatures binaires: détection de format

Fonctionnalités:
• Interface graphique intuitive
• Traitement par lot
• Export JSON/CSV/HTML
• Analyse forensique détaillée

Dépendances optionnelles:
• Pillow: extraction EXIF d'images
• PyPDF2: métadonnées PDF
• Mutagen: tags audio/vidéo
• python-docx: documents Word

Auteur: Assistant IA
Date: Juillet 2025
Licence: MIT

⚠️ Usage légal uniquement:
• Vos propres fichiers
• Autorisation du propriétaire
• Conformité aux lois locales
        """
        
        text_widget = scrolledtext.ScrolledText(self.about_frame, wrap='word')
        text_widget.pack(fill='both', expand=True, padx=10, pady=10)
        text_widget.insert('1.0', about_text)
        text_widget.configure(state='disabled')
    
    def add_files(self):
        """Ajoute des fichiers à analyser."""
        files = filedialog.askopenfilenames(
            title="Sélectionner des fichiers à analyser",
            filetypes=[("Tous les fichiers", "*.*")]
        )
        
        for file_path in files:
            if file_path not in self.selected_files:
                self.selected_files.append(file_path)
                self.files_listbox.insert(tk.END, os.path.basename(file_path))
    
    def add_directory(self):
        """Ajoute tous les fichiers d'un dossier."""
        directory = filedialog.askdirectory(title="Sélectionner un dossier")
        
        if directory:
            added_count = 0
            for root, dirs, files in os.walk(directory):
                for file in files:
                    file_path = os.path.join(root, file)
                    if file_path not in self.selected_files:
                        self.selected_files.append(file_path)
                        self.files_listbox.insert(tk.END, os.path.relpath(file_path, directory))
                        added_count += 1
            
            messagebox.showinfo("Dossier ajouté", f"{added_count} fichiers ajoutés")
    
    def clear_files(self):
        """Efface la liste des fichiers."""
        self.selected_files.clear()
        self.files_listbox.delete(0, tk.END)
    
    def start_extraction(self):
        """Démarre l'extraction des métadonnées."""
        if not self.selected_files:
            messagebox.showerror("Erreur", "Veuillez sélectionner des fichiers à analyser")
            return
        
        self.extract_button.configure(state='disabled')
        self.progress_var.set(0)
        self.status_var.set("Extraction en cours...")
        
        # Démarrer dans un thread séparé
        threading.Thread(target=self.run_extraction, daemon=True).start()
    
    def run_extraction(self):
        """Exécute l'extraction dans un thread séparé."""
        try:
            def progress_callback(progress, message):
                self.root.after(0, lambda: self.progress_var.set(progress))
                self.root.after(0, lambda: self.status_var.set(message))
            
            self.results = self.extractor.extract_batch(
                self.selected_files, 
                lambda p, m: progress_callback(p, m)
            )
            
            self.root.after(0, self.extraction_complete)
            
        except Exception as e:
            self.root.after(0, lambda: messagebox.showerror("Erreur", f"Erreur durant l'extraction: {e}"))
            self.root.after(0, lambda: self.extract_button.configure(state='normal'))
    
    def extraction_complete(self):
        """Appelé quand l'extraction est terminée."""
        self.extract_button.configure(state='normal')
        self.progress_var.set(100)
        self.status_var.set(f"Extraction terminée: {len(self.results)} fichiers traités")
        
        success_count = len([r for r in self.results if 'error' not in r])
        messagebox.showinfo("Extraction terminée", 
                          f"Métadonnées extraites!\n{success_count}/{len(self.results)} fichiers traités avec succès")
        
        self.refresh_results()
    
    def refresh_results(self):
        """Actualise l'affichage des résultats."""
        self.results_text.delete('1.0', tk.END)
        
        if not self.results:
            self.results_text.insert(tk.END, "Aucun résultat disponible. Lancez d'abord une extraction.")
            return
        
        for i, result in enumerate(self.results, 1):
            filename = result.get('file_info', {}).get('filename', 'Fichier inconnu')
            self.results_text.insert(tk.END, f"\n{'='*60}\n")
            self.results_text.insert(tk.END, f"📁 FICHIER {i}: {filename}\n")
            self.results_text.insert(tk.END, f"{'='*60}\n")
            
            if 'error' in result:
                self.results_text.insert(tk.END, f"❌ ERREUR: {result['error']}\n")
            else:
                # Afficher les métadonnées de façon structurée
                self.results_text.insert(tk.END, json.dumps(result, indent=2, ensure_ascii=False, default=str))
                self.results_text.insert(tk.END, "\n\n")
        
        self.results_text.see('1.0')
    
    def save_results(self):
        """Sauvegarde les résultats."""
        if not self.results:
            messagebox.showwarning("Aucun résultat", "Aucun résultat à sauvegarder")
            return
        
        format_ext = {
            'json': '.json',
            'csv': '.csv',
            'html': '.html'
        }
        
        output_format = self.output_format_var.get()
        filename = filedialog.asksaveasfilename(
            title="Sauvegarder les résultats",
            defaultextension=format_ext[output_format],
            filetypes=[(f"Fichiers {output_format.upper()}", f"*{format_ext[output_format]}")]
        )
        
        if filename:
            try:
                self.extractor.export_results(self.results, output_format, filename)
                messagebox.showinfo("Sauvegarde réussie", f"Résultats sauvegardés dans {filename}")
            except Exception as e:
                messagebox.showerror("Erreur de sauvegarde", str(e))
    
    def clear_results(self):
        """Efface les résultats."""
        if messagebox.askyesno("Confirmation", "Effacer tous les résultats?"):
            self.results.clear()
            self.results_text.delete('1.0', tk.END)
            self.status_var.set("Résultats effacés")
    
    def run(self):
        """Lance l'application."""
        self.root.mainloop()


def main():
    """Fonction principale avec interface en ligne de commande."""
    parser = argparse.ArgumentParser(
        description="Extracteur de Métadonnées Forensique",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Exemples d'utilisation:
  python metadata_extractor.py file.jpg
  python metadata_extractor.py *.pdf --format json --output report.json
  python metadata_extractor.py folder/ --recursive --no-hash
  python metadata_extractor.py --gui
        """
    )
    
    parser.add_argument('files', nargs='*', help='Fichiers à analyser')
    parser.add_argument('--gui', action='store_true', help='Lancer l\'interface graphique')
    parser.add_argument('--recursive', '-r', action='store_true', help='Analyse récursive des dossiers')
    parser.add_argument('--format', '-f', choices=['json', 'csv', 'html'], default='json',
                       help='Format de sortie')
    parser.add_argument('--output', '-o', help='Fichier de sortie')
    parser.add_argument('--no-hash', action='store_true', help='Ne pas calculer les hachages')
    parser.add_argument('--verbose', '-v', action='store_true', help='Mode verbose')
    
    args = parser.parse_args()
    
    # Lancer l'interface graphique si demandée
    if args.gui:
        try:
            app = MetadataExtractorGUI()
            app.run()
        except ImportError:
            print("❌ Tkinter non disponible. Interface graphique indisponible.")
        return
    
    # Interface en ligne de commande
    if not args.files:
        print("❌ Veuillez spécifier des fichiers à analyser ou utiliser --gui")
        parser.print_help()
        return
    
    print("🔍 Extracteur de Métadonnées v1.0")
    print("=" * 50)
    
    extractor = MetadataExtractor()
    
    # Construire la liste des fichiers
    file_list = []
    for file_pattern in args.files:
        if os.path.isfile(file_pattern):
            file_list.append(file_pattern)
        elif os.path.isdir(file_pattern):
            if args.recursive:
                for root, dirs, files in os.walk(file_pattern):
                    file_list.extend([os.path.join(root, f) for f in files])
            else:
                file_list.extend([os.path.join(file_pattern, f) 
                                for f in os.listdir(file_pattern) 
                                if os.path.isfile(os.path.join(file_pattern, f))])
        else:
            # Pattern avec wildcards
            import glob
            file_list.extend(glob.glob(file_pattern))
    
    if not file_list:
        print("❌ Aucun fichier trouvé.")
        return
    
    print(f"📊 {len(file_list)} fichier(s) à analyser")
    
    # Extraction avec callback de progrès si verbose
    def progress_callback(progress, message):
        if args.verbose:
            print(f"[{progress:3.0f}%] {message}")
    
    results = extractor.extract_batch(
        file_list, 
        progress_callback if args.verbose else None
    )
    
    # Export des résultats
    content = extractor.export_results(results, args.format, args.output)
    
    if not args.output:
        print("\n" + content)
    
    # Statistiques finales
    success_count = len([r for r in results if 'error' not in r])
    print(f"\n✅ Extraction terminée: {success_count}/{len(results)} fichiers traités avec succès")


if __name__ == "__main__":
    main()
